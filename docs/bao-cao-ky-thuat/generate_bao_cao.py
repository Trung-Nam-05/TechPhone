# -*- coding: utf-8 -*-
"""Generate Bao_cao_Ky_thuat_TechPhone.docx — run: py generate_bao_cao.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

OUTPUT = Path(__file__).resolve().parent / "Bao_cao_Ky_thuat_TechPhone.docx"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_center_title(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)


def add_image_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Chèn ảnh minh chứng: {caption}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(12)
    doc.add_paragraph()


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_font(doc)

    # --- Bìa ---
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "BÁO CÁO KỸ THUẬT ỨNG DỤNG TMĐT", 18)
    add_center_title(doc, "TECHPHONE — CÁC TOPIC CÔNG NGHỆ ĐÃ ÁP DỤNG", 14)
    doc.add_paragraph()
    add_center_title(doc, "Môn: Thương mại điện tử", 13, bold=False)
    add_center_title(doc, "Sinh viên: ................................", 13, bold=False)
    add_center_title(doc, "MSSV: ................................", 13, bold=False)
    add_center_title(doc, "Lớp: ................................", 13, bold=False)
    add_center_title(doc, "Giảng viên: ................................", 13, bold=False)
    doc.add_paragraph()
    add_center_title(doc, "Năm học 2025 – 2026", 13, bold=False)
    doc.add_page_break()

    # --- Giới thiệu ---
    add_heading(doc, "Giới thiệu", 1)
    add_body(
        doc,
        "Báo cáo trình bày các topic kỹ thuật mà giảng viên yêu cầu, được ánh xạ sang cách em đã triển khai "
        "thực tế trên dự án TechPhone (React + Node.js + MongoDB). Phương pháp trình bày: mô tả công nghệ em "
        "đã chọn, lý do chọn, luồng hoạt động trong app, file mã nguồn liên quan và minh chứng chạy được — "
        "không chỉ lý thuyết chung.",
    )
    add_body(
        doc,
        "Trọng tâm báo cáo: (1) Chat realtime Socket.io + chatbot AI, (2) Cổng thanh toán VNPAY, "
        "(3) Tích hợp vận chuyển GHN. Các topic khác (Pusher, Notification Hub, RAG, Mining Data…) "
        "được map sang giải pháp tương đương hoặc ghi rõ phạm vi đã/chưa triển khai.",
    )

    add_heading(doc, "Bảng map topic giảng viên ↔ TechPhone", 2)
    add_table(
        doc,
        [
            ("Pusher realtime", "Socket.io 4", "Chat nhân viên + admin inbox — tự host"),
            ("Notification hub", "Unread badge + Socket events", "FAB widget, tab Nhân viên, sidebar admin"),
            ("Room live", "Socket.io rooms", "conversation:{id}, admin:support"),
            ("Chatbot RAG", "Gemini + Function Calling", "Tra DB sản phẩm/đơn — không vector RAG"),
            ("Mining Data → gợi ý", "Analytics + funnel", "Admin dashboard top products; chưa ML recommend"),
            ("Bảo mật", "PBKDF2, JWT, HMAC VNPAY", "server/src/utils/auth.js, vnpay.js"),
            ("Deploy", "Monorepo npm build", "Dev local; production hướng dẫn VPS/ngrok"),
            ("VNPAY", "Đã tích hợp đầy đủ", "Return URL + IPN, auto confirm đơn"),
            ("GHN", "Đã tích hợp DEV", "Admin confirm → tạo vận đơn → demo progress"),
        ],
        headers=("Topic yêu cầu", "TechPhone áp dụng", "Ghi chú"),
    )

    doc.add_page_break()

    # ===== TOPIC 1: REALTIME (SOCKET.IO) =====
    add_heading(doc, "Topic 1 — Realtime: Socket.io (thay Pusher)", 1)
    add_body(
        doc,
        "Giảng viên đề cập Pusher — dịch vụ realtime SaaS. TechPhone dùng Socket.io (open-source) gắn trực tiếp "
        "vào HTTP server Express, tự kiểm soát dữ liệu chat, không phụ thuộc bên thứ ba. Chức năng tương đương: "
        "push tin nhắn tức thì hai chiều giữa khách và admin.",
    )

    add_heading(doc, "1.1. Kiến trúc đã áp dụng", 2)
    add_table(
        doc,
        [
            ("server/src/socket.js", "Khởi tạo Socket.io trên cùng port 4000 với Express"),
            ("server/src/services/supportChat.js", "Lưu Message/Conversation MongoDB, unread count"),
            ("server/src/routes/support.chat.js", "REST fallback: GET/POST messages"),
            ("server/src/routes/admin.support.js", "Admin inbox API"),
            ("src/context/SupportChatContext.jsx", "Client: kết nối socket, state tin nhắn"),
            ("src/components/HumanSupportPanel.jsx", "UI chat tab Nhân viên"),
            ("src/pages/AdminSupport.jsx", "Admin trả lời khách realtime"),
        ],
        headers=("File", "Vai trò"),
    )

    add_heading(doc, "1.2. Luồng hoạt động step-by-step", 2)
    add_bullet(doc, "Bước 1 — Khách đăng nhập, mở widget chat → tab Nhân viên.")
    add_bullet(doc, "Bước 2 — POST /api/support/conversations tạo/mở Conversation (MongoDB).")
    add_bullet(doc, "Bước 3 — Client connect Socket.io với JWT: io(SOCKET_URL, { auth: { token } }).")
    add_bullet(doc, "Bước 4 — Emit conversation:join → server join room conversation:{id}.")
    add_bullet(doc, "Bước 5 — Khách gửi tin: message:send → supportChat.sendMessage() → lưu DB.")
    add_bullet(doc, "Bước 6 — Server broadcast message:new tới room + admin:support.")
    add_bullet(doc, "Bước 7 — Admin tại /admin/support nhận tin, trả lời → khách thấy ngay + badge unread.")

    add_heading(doc, "1.3. Xác thực Socket", 2)
    add_body(
        doc,
        "Middleware io.use() đọc handshake.auth.token → verifyAccessToken (JWT HS256) → gắn socket.data.userId, "
        "socket.data.role. Không có token hợp lệ → từ chối kết nối. Admin role tự join room admin:support.",
    )

    add_heading(doc, "1.4. Cấu hình .env", 2)
    add_code_block(
        doc,
        "SOCKET_CORS_ORIGIN=http://localhost:5173\nVITE_SOCKET_URL=http://localhost:4000",
    )
    add_image_placeholder(doc, "Widget chat tab Nhân viên — khách gửi tin")
    add_image_placeholder(doc, "Admin /admin/support — inbox + trả lời realtime")

    doc.add_page_break()

    # ===== TOPIC 2: NOTIFICATION HUB =====
    add_heading(doc, "Topic 2 — Notification Hub (thông báo trong app)", 1)
    add_body(
        doc,
        "TechPhone chưa dùng Firebase FCM hay service Notification Hub riêng. Thay vào đó, hệ thống thông báo "
        "realtime nội bộ qua Socket.io events + đếm unread lưu trên Conversation model.",
    )
    add_table(
        doc,
        [
            ("unreadByCustomer", "Số tin admin gửi mà khách chưa đọc"),
            ("unreadByAdmin", "Số tin khách gửi mà admin chưa đọc"),
            ("message:new", "Push tin mới tới client đang online"),
            ("conversation:updated", "Cập nhật preview + sort inbox"),
            ("FAB badge ChatWidget", "Hiển thị tổng unread trên nút chat"),
            ("AdminLayout sidebar", "Badge tổng adminUnreadTotal trên menu Hỗ trợ"),
        ],
        headers=("Cơ chế", "Mô tả"),
    )
    add_body(
        doc,
        "Khi user offline: lần mở widget tiếp theo REST GET /api/support/conversations/me load unread từ DB. "
        "Polling fallback mỗi vài giây nếu socket disconnect.",
    )
    add_image_placeholder(doc, "Badge unread trên widget chat và sidebar admin")

    # ===== TOPIC 3: ROOM LIVE =====
    add_heading(doc, "Topic 3 — Room live (Socket.io rooms)", 1)
    add_body(
        doc,
        "Socket.io Room là nhóm socket subscribe cùng channel — chỉ thành viên room nhận event. TechPhone dùng 2 loại room:",
    )
    add_table(
        doc,
        [
            ("conversation:{conversationId}", "Khách + admin đang xem cùng cuộc chat"),
            ("admin:support", "Mọi admin online — nhận tin mới từ mọi khách"),
        ],
        headers=("Room", "Thành viên / Mục đích"),
    )
    add_body(doc, "Events trong room:")
    add_bullet(doc, "conversation:join — client vào room sau khi verify quyền")
    add_bullet(doc, "message:new, message:delivered — tin nhắn realtime")
    add_bullet(doc, "typing:start / typing:stop — indicator đang gõ")
    add_bullet(doc, "conversation:read — reset unread khi đọc hết")

    doc.add_page_break()

    # ===== TOPIC 4: CHATBOT =====
    add_heading(doc, "Topic 4 — Chatbot AI (map Chatbot RAG)", 1)
    add_body(
        doc,
        "Topic RAG (Retrieval-Augmented Generation) thường dùng vector DB + embed document. TechPhone chọn "
        "Google Gemini với Function Calling — AI gọi tool tra MongoDB trực tiếp (searchProducts, getMyOrders…). "
        "Ưu điểm: giá/tồn kho luôn realtime; không cần re-index CSV Q&A.",
    )
    add_table(
        doc,
        [
            ("server/src/services/gemini.js", "System prompt, model fallback, tool loop"),
            ("server/src/services/aiChatTools.js", "5 tools tra DB"),
            ("server/src/services/aiChat.js", "Session, rate limit 20/phút"),
            ("server/src/routes/ai.chat.js", "GET session, POST message, DELETE session"),
            ("src/components/AiChatPanel.jsx", "UI tab Trợ lý AI"),
            ("@google/generative-ai", "SDK gọi Gemini API"),
        ],
        headers=("File / Thư viện", "Vai trò"),
    )
    add_heading(doc, "4.1. Luồng chatbot", 2)
    add_bullet(doc, "Khách gõ câu hỏi → POST /api/ai-chat")
    add_bullet(doc, "Gemini phân tích → gọi tool (vd: searchProducts('iPhone'))")
    add_bullet(doc, "Tool query MongoDB → trả JSON sản phẩm")
    add_bullet(doc, "Gemini tổng hợp câu trả lời ngắn → lưu AiMessage")
    add_bullet(doc, "Phức tạp → gợi ý chuyển tab Nhân viên (HumanSupportPanel)")
    add_code_block(doc, "GEMINI_API_KEY=...\nGEMINI_MODEL=gemini-flash-latest\nAI_CHAT_RATE_LIMIT=20")
    add_image_placeholder(doc, "Tab Trợ lý AI — hỏi giá sản phẩm, AI tra DB trả lời")

    doc.add_page_break()

    # ===== TOPIC 5: VNPAY =====
    add_heading(doc, "Topic 5 — Cổng thanh toán VNPAY (trọng tâm)", 1)
    add_body(
        doc,
        "VNPAY là cổng thanh toán phổ biến tại Việt Nam (thẻ ATM, ví, QR). TechPhone tích hợp sandbox VNPAY "
        "với luồng redirect + xác minh chữ ký HMAC-SHA512 + cập nhật đơn tự động.",
    )

    add_heading(doc, "5.1. File mã nguồn", 2)
    add_table(
        doc,
        [
            ("server/src/services/vnpay.js", "buildVnpayPaymentUrl, verifyVnpayCallback, sign HMAC"),
            ("server/src/routes/payments.vnpay.js", "GET /return, GET /ipn"),
            ("server/src/routes/orders.js", "Tạo đơn vnpay → trả paymentUrl"),
            ("src/pages/Checkout.jsx", "Chọn VNPAY → redirect paymentUrl"),
            ("src/pages/VnpayResult.jsx", "Landing sau redirect"),
        ],
        headers=("File", "Vai trò"),
    )

    add_heading(doc, "5.2. Cấu hình step-by-step", 2)
    add_bullet(doc, "Bước 1 — Đăng ký sandbox: https://sandbox.vnpayment.vn/devreg/")
    add_bullet(doc, "Bước 2 — Lấy VNPAY_TMN_CODE và VNPAY_HASH_SECRET → ghi .env")
    add_bullet(doc, "Bước 3 — Dev local: chạy ngrok http 4000 → set API_PUBLIC_URL=https://xxx.ngrok-free.app")
    add_bullet(doc, "Bước 4 — VNPAY cần gọi callback về server public — localhost không được")
    add_bullet(doc, "Bước 5 — CLIENT_ORIGIN=http://localhost:5173 để redirect khách về React sau thanh toán")

    add_heading(doc, "5.3. Luồng thanh toán trong app", 2)
    add_bullet(doc, "1. Checkout chọn VNPAY → POST /api/orders { paymentMethod: 'vnpay' }")
    add_bullet(doc, "2. Order tạo status=pending, paymentStatus=pending")
    add_bullet(doc, "3. buildVnpayPaymentUrl: sort params → HMAC-SHA512 → URL sandbox.vnpayment.vn")
    add_bullet(doc, "4. vnp_ReturnUrl = {API_PUBLIC_URL}/api/payments/vnpay/return")
    add_bullet(doc, "5. vnp_IpnUrl = {API_PUBLIC_URL}/api/payments/vnpay/ipn (chỉ khi HTTPS public)")
    add_bullet(doc, "6. Khách thanh toán OK (responseCode=00)")
    add_bullet(doc, "7. IPN/Return: verifyVnpayCallback → kiểm tra amount → paymentStatus=paid, status=confirmed")
    add_bullet(doc, "8. Redirect → /checkout/vnpay-result?success=1&orderId=...")

    add_heading(doc, "5.4. Thuật toán bảo mật VNPAY", 2)
    add_body(
        doc,
        "Ký: sort key alphabet → URLSearchParams encode → HMAC-SHA512(hashSecret, signData). "
        "Verify callback: loại vnp_SecureHash, tính lại hash, so sánh. Kiểm tra vnp_Amount === order.total * 100 "
        "chống giả mạo số tiền.",
    )
    add_code_block(
        doc,
        "VNPAY_TMN_CODE=...\nVNPAY_HASH_SECRET=...\nVNPAY_PAYMENT_URL=https://sandbox.vnpayment.vn/paymentv2/vpcpay.html\nAPI_PUBLIC_URL=https://xxx.ngrok-free.app\nCLIENT_ORIGIN=http://localhost:5173",
    )
    add_image_placeholder(doc, "Checkout — chọn VNPAY")
    add_image_placeholder(doc, "Trang sandbox VNPAY thanh toán")
    add_image_placeholder(doc, "VnpayResult success + Account orders paymentStatus=paid")

    doc.add_page_break()

    # ===== TOPIC 6: GHN =====
    add_heading(doc, "Topic 6 — Tích hợp vận chuyển GHN (trọng tâm)", 1)
    add_body(
        doc,
        "GHN (Giao Hàng Nhanh) cung cấp API tạo vận đơn, tra cứu trạng thái, master data địa chỉ. "
        "TechPhone dùng môi trường DEV (dev-online-gateway.ghn.vn) — không có shipper thật; "
        "admin bấm xác nhận thủ công rồi job demo tự chuyển trạng thái đơn.",
    )

    add_heading(doc, "6.1. File mã nguồn", 2)
    add_table(
        doc,
        [
            ("server/src/services/ghn.js", "HTTP client GHN API, submitOrder, getOrderDetail"),
            ("server/src/services/ghnAddress.js", "Resolve province/district/ward → GHN IDs"),
            ("server/src/services/ghnShipment.js", "createGhnShipmentForOrder, sync status"),
            ("server/src/services/ghnSync.js", "Job poll trạng thái (GHN_POLL_ENABLED)"),
            ("server/src/services/ghnDemoProgress.js", "Job demo chuyển trạng thái DEV"),
            ("server/src/routes/shipping.ghn.js", "API provinces/districts/wards cho checkout"),
            ("server/src/routes/admin.orders.js", "POST confirm-fulfillment → gửi GHN"),
            ("src/pages/AdminOrders.jsx", "Nút Xác nhận & gửi GHN"),
            ("src/pages/OrderDetail.jsx", "Stepper tiến trình (khách không thấy mã vận đơn)"),
        ],
        headers=("File", "Vai trò"),
    )

    add_heading(doc, "6.2. Cấu hình step-by-step", 2)
    add_bullet(doc, "Bước 1 — Đăng ký tài khoản GHN DEV, lấy token + shop ID")
    add_bullet(doc, "Bước 2 — .env: GHN_ENABLED=true, GHN_API_TOKEN, GHN_SHOP_ID, GHN_API_URL=dev-online-gateway")
    add_bullet(doc, "Bước 3 — GHN_DEMO_PROGRESS_ENABLED=true — tự chuyển await_pickup → shipping → completed")
    add_bullet(doc, "Bước 4 — npm run ghn:test kiểm tra kết nối API")

    add_heading(doc, "6.3. Luồng vận chuyển trong app", 2)
    add_bullet(doc, "1. Đơn COD/VNPAY paid → status=confirmed")
    add_bullet(doc, "2. Admin /admin/orders → tìm đơn → Xác nhận & gửi GHN")
    add_bullet(doc, "3. POST /api/admin/orders/:id/confirm-fulfillment")
    add_bullet(doc, "4. resolveGhnAddress(shippingInfo) → districtId, wardCode")
    add_bullet(doc, "5. submitOrder() gọi GHN API → nhận orderCode (labelId)")
    add_bullet(doc, "6. Order.shipment.labelId lưu DB; status → await_pickup")
    add_bullet(doc, "7. ghnDemoProgress job (DEV) tự chuyển picked → shipping → completed theo timeout .env")
    add_bullet(doc, "8. Khách xem stepper tại /account/orders/:id — orderSanitize ẩn mã vận đơn GHN")

    add_heading(doc, "6.4. State machine trạng thái đơn", 2)
    add_code_block(
        doc,
        "pending → confirmed → await_pickup → picked → shipping → completed\n(cancelled | delivery_failed | returned)",
    )
    add_code_block(
        doc,
        "GHN_ENABLED=true\nGHN_API_TOKEN=...\nGHN_SHOP_ID=...\nGHN_DEMO_PROGRESS_ENABLED=true\nGHN_DEMO_PROGRESS_MS=5000",
    )
    add_image_placeholder(doc, "Admin Orders — nút Xác nhận & gửi GHN")
    add_image_placeholder(doc, "OrderDetail — stepper vận chuyển khách hàng")

    doc.add_page_break()

    # ===== TOPIC 7: MINING DATA =====
    add_heading(doc, "Topic 7 — Mining Data & hệ gợi ý", 1)
    add_body(
        doc,
        "TechPhone thu thập sự kiện hành vi qua AnalyticsEvent (POST /api/analytics/events): view_product, "
        "add_to_cart, begin_checkout, purchase. Admin dashboard (/admin/analytics) hiển thị funnel và top products.",
    )
    add_table(
        doc,
        [
            ("AnalyticsEvent model", "eventName, sessionId, userId, metadata, path"),
            ("GET /api/analytics/funnel", "Conversion funnel 4 bước"),
            ("GET /api/analytics/dashboard", "KPI + topProducts bán chạy"),
            ("GET /api/analytics/revenue", "Doanh thu theo thời gian"),
            ("AI getTopProducts", "Gợi ý SP rẻ/đắt theo category — rule-based, chưa ML"),
        ],
        headers=("Thành phần", "Mô tả"),
    )
    add_body(
        doc,
        "Hạn chế: chưa có recommendation engine collaborative filtering. Gợi ý hiện tại dựa trên ranking "
        "doanh số admin + AI tool sort giá. Hướng phát triển: train model từ AnalyticsEvent + Order.items.",
    )
    add_image_placeholder(doc, "Admin Analytics — funnel + top products chart")

    # ===== TOPIC 8: SECURITY =====
    add_heading(doc, "Topic 8 — Bảo mật", 1)
    add_table(
        doc,
        [
            ("Mật khẩu", "PBKDF2-SHA512, 100000 iterations, salt random, timingSafeEqual"),
            ("Session API", "JWT HS256, TTL 7 ngày, Bearer token"),
            ("Login", "Brute-force: max 6 lần/10 phút, khóa 5 phút"),
            ("VNPAY", "HMAC-SHA512 verify callback, kiểm tra amount"),
            ("Socket", "JWT handshake auth"),
            ("Order API", "orderSanitize — ẩn shipment với khách"),
            ("Admin", "requireAuth + requireAdmin middleware"),
            ("Ownership", "User chỉ xem đơn/giỏ của mình; guest dùng x-session-id"),
        ],
        headers=("Lớp", "Biện pháp"),
    )

    # ===== TOPIC 9: DEPLOY =====
    add_heading(doc, "Topic 9 — Deploy & vận hành", 1)
    add_body(doc, "Hiện trạng: phát triển local. Quy trình deploy đề xuất:")
    add_bullet(doc, "Frontend: npm run build → thư mục dist/ → host static (Vercel, Nginx)")
    add_bullet(doc, "Backend: node server/src/index.js trên VPS (PM2) — port 4000")
    add_bullet(doc, "Database: MongoDB Atlas (đã dùng cloud)")
    add_bullet(doc, "Biến production: MONGODB_URI, JWT_SECRET, API_PUBLIC_URL domain thật, VNPAY production keys")
    add_bullet(doc, "GHN: chuyển GHN_API_URL production khi go-live; tắt GHN_DEMO_PROGRESS")
    add_bullet(doc, "HTTPS bắt buộc cho VNPAY IPN và webhook")

    add_heading(doc, "Topic bổ sung — MongoDB Transaction chống oversell", 2)
    add_body(
        doc,
        "POST /api/orders dùng mongoose.withTransaction: Product.findOneAndUpdate atomic $inc stock, "
        "FlashSale quota, tạo Order — tránh bán vượt tồn kho khi nhiều user checkout đồng thời.",
    )

    add_heading(doc, "Topic bổ sung — Hóa đơn điện tử", 2)
    add_body(
        doc,
        "Toggle checkout invoiceRequested → ElectronicInvoice component trên OrderSuccessResult và OrderDetail. "
        "Demo hóa đơn TMĐT — chưa tích hợp cổng HĐĐT nhà nước.",
    )

    doc.add_page_break()

    # ===== KẾT LUẬN =====
    add_heading(doc, "Kết luận", 1)
    add_body(
        doc,
        "TechPhone đã áp dụng các topic kỹ thuật TMĐT với trọng tâm: Socket.io realtime chat (map Pusher), "
        "notification unread hub nội bộ, room-based messaging, chatbot Gemini function calling (map RAG), "
        "tích hợp VNPAY đầy đủ (Return + IPN + HMAC), tích hợp GHN DEV (tạo vận đơn + demo progress), "
        "analytics mining cơ bản, bảo mật PBKDF2/JWT/HMAC và hướng dẫn deploy. "
        "Em minh chứng bằng mã nguồn có thể chạy (npm run dev:full) và ảnh chụp màn hình đính kèm.",
    )

    add_heading(doc, "Phụ lục — Checklist ảnh minh chứng (16 ảnh)", 2)
    checklist = [
        "Widget chat FAB + tab AI",
        "Tab Nhân viên — chat khách",
        "Admin inbox /admin/support",
        "Badge unread widget + sidebar",
        "Typing indicator (optional)",
        "AI hỏi sản phẩm — tool tra DB",
        "Checkout chọn VNPAY",
        "Trang sandbox VNPAY",
        "VnpayResult success",
        "Account orders — paid",
        "Admin Orders — confirm GHN",
        "Order stepper khách",
        "Admin Analytics funnel",
        ".env che secret (VNPAY, GHN, Gemini)",
        "Network tab — Socket.io WS connection",
        "POST /api/orders response paymentUrl",
    ]
    for i, item in enumerate(checklist, 1):
        add_bullet(doc, f"Ảnh {i}: {item}")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
