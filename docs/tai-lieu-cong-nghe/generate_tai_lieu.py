# -*- coding: utf-8 -*-
"""Generate Tai_lieu_Cong_nghe_TechPhone.docx — run: py generate_tai_lieu.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

OUTPUT = Path(__file__).resolve().parent / "Tai_lieu_Cong_nghe_TechPhone.docx"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_center_title(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_font(doc)

    # --- Trang bìa ---
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "TÀI LIỆU CÔNG NGHỆ DỰ ÁN", 18)
    add_center_title(doc, "TECHPHONE — WEBSITE THƯƠNG MẠI ĐIỆN TỬ", 14)
    doc.add_paragraph()
    add_center_title(doc, "Môn: Thương mại điện tử", 13, bold=False)
    add_center_title(doc, "Sinh viên: ................................", 13, bold=False)
    add_center_title(doc, "MSSV: ................................", 13, bold=False)
    add_center_title(doc, "Lớp: ................................", 13, bold=False)
    add_center_title(doc, "Giảng viên: ................................", 13, bold=False)
    doc.add_paragraph()
    add_center_title(doc, "Năm học 2025 – 2026", 13, bold=False)
    doc.add_page_break()

    # --- Mục lục (text) ---
    add_heading(doc, "Mục lục", 1)
    toc = [
        "1. Giới thiệu tổng quan",
        "2. Kiến trúc hệ thống",
        "3. Công nghệ Frontend",
        "4. Công nghệ Backend",
        "5. Cơ sở dữ liệu MongoDB",
        "6. Tích hợp bên thứ ba",
        "7. Bảo mật và thuật toán mật mã",
        "8. Thuật toán nghiệp vụ TMĐT",
        "9. AI Chatbot và Chat realtime",
        "10. Module chức năng chính",
        "11. Công cụ phát triển và triển khai",
        "12. Phụ lục",
    ]
    for item in toc:
        add_bullet(doc, item)
    doc.add_page_break()

    # --- 1. Giới thiệu ---
    add_heading(doc, "1. Giới thiệu tổng quan", 1)
    add_body(
        doc,
        "TechPhone là website thương mại điện tử (TMĐT) bán điện thoại, laptop, tablet và phụ kiện công nghệ. "
        "Dự án áp dụng mô hình Product Selling — bán hàng hóa vật lý với giá cố định, giỏ hàng, checkout trực tuyến, "
        "thanh toán đa phương thức và vận chuyển tích hợp. Hệ thống được xây dựng theo kiến trúc monorepo: "
        "frontend React + Vite và backend Node.js + Express + MongoDB trong cùng một repository.",
    )
    add_table(
        doc,
        [
            ("Frontend", "React 19, Vite 8, React Router 7", "http://localhost:5173"),
            ("Backend API", "Express 5, Node.js ES Modules", "http://localhost:4000"),
            ("Database", "MongoDB Atlas, Mongoose 8", "Cloud"),
            ("Realtime", "Socket.io 4", "Cùng port API"),
            ("AI", "Google Gemini API", "Cloud"),
        ],
        headers=("Tầng", "Công nghệ", "Vị trí"),
    )
    add_body(
        doc,
        "Repository GitHub: https://github.com/Trung-Nam-05/TechPhone. "
        "Chạy full stack: npm run dev:full (concurrently frontend + backend).",
    )

    # --- 2. Kiến trúc ---
    add_heading(doc, "2. Kiến trúc hệ thống", 1)
    add_body(
        doc,
        "Hệ thống theo mô hình client–server ba tầng: Presentation (React SPA), Application (Express REST + Socket.io), "
        "Data (MongoDB Atlas). Frontend giao tiếp backend qua HTTP JSON; chat nhân viên dùng WebSocket song song.",
    )
    add_body(doc, "Luồng dữ liệu chính:")
    add_bullet(doc, "Khách duyệt SP → GET /api/products → MongoDB")
    add_bullet(doc, "Thêm giỏ → PUT /api/cart (header x-session-id hoặc JWT)")
    add_bullet(doc, "Checkout → POST /api/orders (MongoDB transaction, trừ stock atomic)")
    add_bullet(doc, "VNPAY → redirect cổng → callback Return/IPN → cập nhật paymentStatus")
    add_bullet(doc, "Admin confirm → POST /api/admin/orders/:id/confirm-fulfillment → GHN API")
    add_bullet(doc, "Chat AI → POST /api/ai-chat → Gemini + function calling")
    add_bullet(doc, "Chat nhân viên → Socket.io ↔ supportChat.js ↔ MongoDB")
    add_body(
        doc,
        "[Chèn sơ đồ kiến trúc: Client (React) ↔ Express API ↔ MongoDB / Gemini / VNPAY / GHN]",
    )

    doc.add_page_break()

    # --- 3. Frontend ---
    add_heading(doc, "3. Công nghệ Frontend", 1)

    add_heading(doc, "3.1. Framework và build tool", 2)
    add_table(
        doc,
        [
            ("React", "19.2.4", "UI component-based, Virtual DOM"),
            ("React DOM", "19.2.4", "Render component lên DOM"),
            ("Vite", "8.0.4", "Dev server HMR, build production nhanh"),
            ("@vitejs/plugin-react", "6.0.1", "Hỗ trợ JSX, Fast Refresh"),
        ],
        headers=("Thư viện", "Phiên bản", "Vai trò"),
    )
    add_body(doc, "Entry point: index.html → src/main.jsx → App.jsx (React Router).")

    add_heading(doc, "3.2. Routing và state management", 2)
    add_table(
        doc,
        [
            ("react-router-dom", "7.14.0", "SPA routing: /products, /admin/*, /account/*"),
            ("React Context API", "Built-in", "Global state — không dùng Redux"),
        ],
        headers=("Thư viện / Pattern", "Phiên bản", "Vai trò"),
    )
    add_body(doc, "Các Context Provider (src/main.jsx):")
    add_table(
        doc,
        [
            ("AuthContext", "JWT, login/logout, phân quyền admin"),
            ("CartContext", "Giỏ hàng đồng bộ API /api/cart"),
            ("SupportChatContext", "Chat nhân viên + Socket.io + unread badge"),
            ("AiChatContext", "Session và tin nhắn AI chat"),
            ("AnalyticsContext", "Gửi sự kiện tracking"),
            ("I18nContext", "Ngôn ngữ, format giá VND"),
        ],
        headers=("Context", "Chức năng"),
    )

    add_heading(doc, "3.3. UI và thư viện phụ trợ", 2)
    add_table(
        doc,
        [
            ("lucide-react", "1.7.0", "Icon (cart, chat, admin dashboard)"),
            ("socket.io-client", "4.8.3", "WebSocket client cho chat realtime"),
            ("CSS thuần", "—", "index.css, component CSS; brand color #CB1C22"),
        ],
        headers=("Thư viện", "Phiên bản", "Vai trò"),
    )

    add_heading(doc, "3.4. Cấu trúc thư mục src/", 2)
    add_table(
        doc,
        [
            ("src/pages/", "Màn hình theo URL (Home, Products, Checkout, Admin*, Account*)"),
            ("src/components/", "Component tái sử dụng (Header, ChatWidget, ProductCard…)"),
            ("src/layouts/", "StoreLayout, AdminLayout, AccountLayout"),
            ("src/context/", "State toàn ứng dụng"),
            ("src/config/", "api.js (base URL), socket.js"),
            ("src/data/", "Metadata tĩnh: categories, policies, coupons"),
            ("src/utils/", "Helper: orderInvoice, supportMessageStatus"),
        ],
        headers=("Thư mục", "Nội dung"),
    )

    doc.add_page_break()

    # --- 4. Backend ---
    add_heading(doc, "4. Công nghệ Backend", 1)

    add_heading(doc, "4.1. Runtime và framework", 2)
    add_table(
        doc,
        [
            ("Node.js", "Runtime", "JavaScript server-side, ES Modules (type: module)"),
            ("Express", "5.1.0", "HTTP router, middleware, JSON API"),
            ("dotenv", "17.2.3", "Biến môi trường .env"),
            ("cors", "2.8.5", "Cross-Origin Resource Sharing cho React :5173"),
            ("mongoose", "8.19.2", "ODM MongoDB — schema, query, transaction"),
        ],
        headers=("Thư viện", "Phiên bản / Loại", "Vai trò"),
    )

    add_heading(doc, "4.2. Module Node.js built-in", 2)
    add_table(
        doc,
        [
            ("node:crypto", "PBKDF2 password, JWT HMAC-SHA256, VNPAY HMAC-SHA512"),
            ("node:http", "HTTP server chung cho Express + Socket.io"),
            ("node:dns", "Ép IPv4 khi kết nối MongoDB Atlas (tránh lỗi Windows)"),
        ],
        headers=("Module", "Sử dụng"),
    )

    add_heading(doc, "4.3. Kiến trúc ba lớp backend", 2)
    add_table(
        doc,
        [
            ("server/src/routes/", "HTTP endpoints — mỏng, validate input, gọi service"),
            ("server/src/services/", "Logic nghiệp vụ: pricing, GHN, Gemini, hủy đơn…"),
            ("server/src/models/", "Schema Mongoose — 15 collections"),
            ("server/src/middleware/", "auth.js — JWT requireAuth, optionalAuth, requireAdmin"),
            ("server/src/utils/", "auth, ownership, cart, audit, orderSanitize"),
            ("server/src/constants/", "orderStatus, cartLimits"),
        ],
        headers=("Thư mục", "Vai trò"),
    )

    add_heading(doc, "4.4. Nhóm API REST chính", 2)
    add_table(
        doc,
        [
            ("/api/auth", "Register, login, profile (JWT)"),
            ("/api/products", "Danh sách, chi tiết, search, reviews"),
            ("/api/cart", "Giỏ hàng guest/user (x-session-id)"),
            ("/api/orders", "Tạo đơn, timeline, hủy đơn"),
            ("/api/coupons", "Preview mã giảm giá"),
            ("/api/payments/vnpay", "Return URL + IPN callback"),
            ("/api/shipping/ghn", "Master data địa chỉ GHN"),
            ("/api/support", "Chat nhân viên (khách hàng)"),
            ("/api/ai-chat", "Chatbot Gemini AI"),
            ("/api/admin/*", "Products, orders, inventory, users, analytics, support"),
            ("/api/analytics", "Events tracking, dashboard, revenue"),
        ],
        headers=("Prefix API", "Chức năng"),
    )

    doc.add_page_break()

    # --- 5. Database ---
    add_heading(doc, "5. Cơ sở dữ liệu MongoDB", 1)
    add_body(
        doc,
        "Database hosted trên MongoDB Atlas (cloud). ODM Mongoose 8 cung cấp schema validation, indexes, "
        "và multi-document transactions cho checkout an toàn.",
    )
    add_heading(doc, "5.1. Collections (15 model)", 2)
    add_table(
        doc,
        [
            ("User", "Tài khoản customer/admin, passwordHash, role"),
            ("Product", "Sản phẩm: name, slug, price, stock, category, text index"),
            ("Cart", "Giỏ hàng theo user hoặc sessionId"),
            ("Order", "Đơn hàng: items snapshot, pricing, payment, shipment, installment"),
            ("OrderEvent", "Audit trail thay đổi trạng thái đơn"),
            ("Coupon", "Mã giảm giá: scope, discountType, thời hạn"),
            ("FlashSale", "Giá flash, quota, soldCount"),
            ("InventoryMovement", "Log nhập/xuất/adjustment/order kho"),
            ("Review", "Đánh giá sản phẩm (rating 1–5)"),
            ("Conversation / Message", "Chat nhân viên realtime"),
            ("AiSession / AiMessage", "Session và lịch sử chat AI"),
            ("ShipmentEvent", "Log webhook GHN/GHTK"),
            ("AnalyticsEvent", "Tracking hành vi người dùng"),
            ("AdminAuditLog", "Log hành động admin"),
        ],
        headers=("Model / Collection", "Mục đích"),
    )

    add_heading(doc, "5.2. Index và tìm kiếm", 2)
    add_bullet(doc, "Product: text index { name, brand } — full-text search MongoDB")
    add_bullet(doc, "Product: compound index category + price + brand")
    add_bullet(doc, "Order: index user/sessionId + createdAt; idempotencyKey unique sparse")

    add_heading(doc, "5.3. MongoDB Transaction — chống oversell", 2)
    add_body(
        doc,
        "POST /api/orders sử dụng mongoose.startSession() + withTransaction(). Trong transaction: "
        "Product.findOneAndUpdate với điều kiện stock >= quantity và $inc: { stock: -quantity } — "
        "chỉ trừ kho khi đủ hàng (atomic). Đồng thời reserve FlashSale quota, tạo Order + OrderEvent, xóa Cart.",
    )

    doc.add_page_break()

    # --- 6. Tích hợp ---
    add_heading(doc, "6. Tích hợp bên thứ ba", 1)
    add_table(
        doc,
        [
            ("Google Gemini", "@google/generative-ai", "AI chatbot, function calling tra DB"),
            ("VNPAY", "REST redirect + HMAC-SHA512", "Thanh toán thẻ ATM / ví điện tử"),
            ("GHN (Giao Hàng Nhanh)", "REST API", "Tạo vận đơn, sync trạng thái, master data địa chỉ"),
            ("GHTK", "REST (legacy)", "Webhook cũ — có thể tắt (GHTK_ENABLED=false)"),
            ("MongoDB Atlas", "Mongoose connection", "Database cloud"),
        ],
        headers=("Dịch vụ", "Giao thức / SDK", "Mục đích"),
    )
    add_body(
        doc,
        "VNPAY yêu cầu API_PUBLIC_URL (ngrok/cloudflare tunnel khi dev) để VNPAY gọi callback Return URL "
        "và IPN server-to-server. CLIENT_ORIGIN dùng redirect browser về React sau thanh toán.",
    )

    doc.add_page_break()

    # --- 7. Bảo mật ---
    add_heading(doc, "7. Bảo mật và thuật toán mật mã", 1)
    add_body(doc, "Triển khai trong server/src/utils/auth.js và server/src/services/vnpay.js.")

    add_heading(doc, "7.1. Hash mật khẩu — PBKDF2-SHA512", 2)
    add_body(doc, "Hàm hashPassword() và verifyPassword() — không lưu plaintext password.")
    add_table(
        doc,
        [
            ("Thuật toán", "PBKDF2 (Password-Based Key Derivation Function 2)"),
            ("Hash function", "SHA-512"),
            ("Iterations", "100,000"),
            ("Salt", "16 bytes random (crypto.randomBytes)"),
            ("Output", "64 bytes hex"),
            ("Format DB", "salt:hash (User.passwordHash)"),
            ("So sánh", "crypto.timingSafeEqual — chống timing attack"),
        ],
        headers=("Tham số", "Giá trị"),
    )

    add_heading(doc, "7.2. JWT Access Token — HS256", 2)
    add_table(
        doc,
        [
            ("Algorithm", "HMAC-SHA256 (JWT alg HS256)"),
            ("Secret", "JWT_SECRET trong .env"),
            ("TTL", "7 ngày (604800 giây)"),
            ("Payload", "sub (userId), role, iat, exp"),
            ("Encoding", "Base64URL — header.payload.signature"),
            ("Verify", "Tính lại HMAC + timingSafeEqual + kiểm tra exp"),
        ],
        headers=("Tham số", "Giá trị"),
    )
    add_body(
        doc,
        "Middleware: requireAuth (bắt buộc token), optionalAuth (có token thì gắn req.auth), "
        "requireAdmin (role === admin). Token gửi qua header Authorization: Bearer <token>.",
    )

    add_heading(doc, "7.3. Chống brute-force đăng nhập", 2)
    add_table(
        doc,
        [
            ("Cửa sổ đếm", "10 phút"),
            ("Số lần sai tối đa", "6 lần"),
            ("Thời gian khóa", "5 phút"),
            ("Lưu trữ", "In-memory Map (theo email key)"),
        ],
        headers=("Rule", "Giá trị"),
    )

    add_heading(doc, "7.4. VNPAY chữ ký HMAC-SHA512", 2)
    add_body(
        doc,
        "buildVnpayPaymentUrl(): sort keys params alphabetically → build URL-encoded query string "
        "(URLSearchParams) → HMAC-SHA512(hashSecret, signData) → vnp_SecureHash. "
        "verifyVnpayCallback(): loại vnp_SecureHash, tính lại hash, so sánh trước khi cập nhật đơn paid.",
    )

    add_heading(doc, "7.5. Ownership giỏ hàng và đơn hàng", 2)
    add_body(
        doc,
        "User đăng nhập: gắn user ObjectId. Guest: header x-session-id (UUID client tạo). "
        "Helper getOwnershipFilter() trong utils/ownership.js — đảm bảo mỗi user chỉ truy cập dữ liệu của mình.",
    )

    doc.add_page_break()

    # --- 8. Thuật toán nghiệp vụ ---
    add_heading(doc, "8. Thuật toán nghiệp vụ TMĐT", 1)

    add_heading(doc, "8.1. Tính giá đơn hàng (pricing.js)", 2)
    add_body(doc, "Công thức:")
    add_bullet(doc, "subtotal = tổng (linePrice × quantity)")
    add_bullet(doc, "productDiscount = giảm từ oldPrice hoặc flash sale price")
    add_bullet(doc, "couponDiscount = tối đa 2 mã (scope product + shipping)")
    add_bullet(doc, "shippingFee = 30.000 VND (DEFAULT_SHIPPING_FEE)")
    add_bullet(doc, "total = subtotal - discounts + shippingFee - shippingDiscount")
    add_body(
        doc,
        "Coupon: percentage hoặc fixed; cap maxDiscountValue; kiểm tra minOrderValue, usageLimit, startsAt/endsAt.",
    )

    add_heading(doc, "8.2. Flash Sale (flashSale.js)", 2)
    add_body(
        doc,
        "Trạng thái: inactive, upcoming, active, ended, sold_out. pickBestFlashSale chọn sale active giá thấp nhất. "
        "Checkout: FlashSale.findOneAndUpdate atomic tăng soldCount nếu còn quota.",
    )

    add_heading(doc, "8.3. State machine đơn hàng (orderStatus.js)", 2)
    add_body(doc, "Vòng đời: pending → confirmed → await_pickup → picked → shipping → completed")
    add_body(
        doc,
        "shouldTransitionOrderStatus(): terminal status (completed, cancelled…) luôn cho phép; "
        "status khác chỉ chuyển tiến (rank cao hơn) — tránh lùi trạng thái khi sync GHN.",
    )

    add_heading(doc, "8.4. Tìm kiếm sản phẩm — scoring relevance (productSearch.js)", 2)
    add_table(
        doc,
        [
            ("Tên exact match", "1000 điểm"),
            ("Tên startsWith", "800"),
            ("Tên includes", "600"),
            ("Brand startsWith", "500"),
            ("Brand includes", "400"),
            ("Slug match", "300"),
        ],
        headers=("Loại match", "Điểm"),
    )

    add_heading(doc, "8.5. Idempotency và COD/VNPAY tự động", 2)
    add_bullet(doc, "Header x-idempotency-key — tránh tạo đơn trùng khi double-click")
    add_bullet(doc, "COD: status=confirmed, paymentStatus=pending (thu khi giao)")
    add_bullet(doc, "VNPAY: status=pending → IPN/Return responseCode=00 → paymentStatus=paid, status=confirmed")

    add_heading(doc, "8.6. Hóa đơn điện tử", 2)
    add_body(
        doc,
        "Toggle checkout invoiceRequested → lưu Order.invoiceRequested + ghi chú shippingInfo.note. "
        "Component ElectronicInvoice hiển thị trên OrderSuccessResult và OrderDetail.",
    )

    doc.add_page_break()

    # --- 9. AI & Chat ---
    add_heading(doc, "9. AI Chatbot và Chat realtime", 1)

    add_heading(doc, "9.1. Google Gemini AI", 2)
    add_table(
        doc,
        [
            ("SDK", "@google/generative-ai 0.24.1"),
            ("Model", "gemini-flash-latest (+ fallback chain)"),
            ("Pattern", "Function calling — tool loop"),
            ("Config", "temperature 0.4, maxOutputTokens 280"),
            ("Rate limit", "AI_CHAT_RATE_LIMIT (mặc định 20 tin/phút)"),
        ],
        headers=("Thành phần", "Chi tiết"),
    )
    add_body(doc, "AI Tools (aiChatTools.js):")
    add_table(
        doc,
        [
            ("searchProducts", "Text search + filter category/brand/price"),
            ("getTopProducts", "Sort giá asc/desc theo category"),
            ("getProductDetail", "Chi tiết SP + stock từ DB"),
            ("getMyOrders", "Đơn hàng user đăng nhập"),
            ("getOrderTimeline", "Tiến trình vận chuyển đơn"),
        ],
        headers=("Tool", "Chức năng"),
    )

    add_heading(doc, "9.2. Socket.io — Chat nhân viên", 2)
    add_table(
        doc,
        [
            ("Server", "server/src/socket.js"),
            ("Client", "socket.io-client + SupportChatContext"),
            ("Auth", "JWT trong handshake.auth.token"),
            ("Rooms", "conversation:{id}, admin:support"),
        ],
        headers=("Thành phần", "Chi tiết"),
    )
    add_body(
        doc,
        "Events: message:send, message:new, typing:update, conversation:read, conversation:updated. "
        "Admin inbox: /admin/support — danh sách khách, badge unread, trả lời realtime.",
    )

    doc.add_page_break()

    # --- 10. Module chức năng ---
    add_heading(doc, "10. Module chức năng chính", 1)
    add_table(
        doc,
        [
            ("Storefront", "Products, ProductDetail, Cart, Checkout, Home, Policies"),
            ("Thanh toán", "COD tự xác nhận, VNPAY redirect+IPN, Trả góp pending review"),
            ("Tài khoản khách", "/account: profile, orders, order detail stepper, security"),
            ("Admin", "Dashboard, products CRUD, orders+GHN, inventory, flash sales, users, analytics"),
            ("Chat", "ChatWidget 2 tab: Trợ lý AI (Gemini) + Nhân viên (Socket.io)"),
            ("Hóa đơn", "ElectronicInvoice trên success screen và order detail"),
            ("Analytics", "POST events, admin dashboard KPI, revenue chart"),
        ],
        headers=("Module", "Tính năng"),
    )

    # --- 11. DevOps ---
    add_heading(doc, "11. Công cụ phát triển và triển khai", 1)
    add_table(
        doc,
        [
            ("npm", "Package manager, scripts"),
            ("concurrently", "Chạy frontend + backend (dev:full)"),
            ("nodemon", "Auto-restart API khi đổi file JS"),
            ("ESLint 9", "Lint code React"),
            ("Git + GitHub", "Version control"),
        ],
        headers=("Công cụ", "Vai trò"),
    )
    add_body(doc, "Scripts npm:")
    add_table(
        doc,
        [
            ("npm run dev:full", "Frontend :5173 + Backend :4000"),
            ("npm run seed", "Nạp ~29 sản phẩm + tài khoản admin"),
            ("npm run build", "Build production frontend (dist/)"),
            ("npm run ghn:test", "Test kết nối GHN API"),
        ],
        headers=("Lệnh", "Mô tả"),
    )

    doc.add_page_break()

    # --- 12. Phụ lục ---
    add_heading(doc, "12. Phụ lục", 1)

    add_heading(doc, "12.1. Biến môi trường (.env)", 2)
    add_table(
        doc,
        [
            ("MONGODB_URI", "Connection string MongoDB Atlas"),
            ("JWT_SECRET", "Secret ký JWT HS256"),
            ("VITE_API_BASE_URL", "Frontend gọi API (http://localhost:4000)"),
            ("CLIENT_ORIGIN", "CORS + redirect VNPAY về React"),
            ("API_PUBLIC_URL", "VNPAY Return/IPN callback (ngrok khi dev)"),
            ("VNPAY_TMN_CODE, VNPAY_HASH_SECRET", "Cấu hình cổng VNPAY"),
            ("GHN_ENABLED, GHN_API_TOKEN, GHN_SHOP_ID", "Vận chuyển GHN"),
            ("GEMINI_API_KEY, GEMINI_MODEL", "AI chatbot"),
            ("ADMIN_EMAIL, ADMIN_PASSWORD", "Tài khoản admin khi seed"),
        ],
        headers=("Biến", "Mục đích"),
    )

    add_heading(doc, "12.2. Mô hình kinh doanh", 2)
    add_body(
        doc,
        "TechPhone áp dụng Product Selling — bán sản phẩm vật lý (điện thoại, laptop, phụ kiện). "
        "Không áp dụng Subscription, Usage-Based, Transaction Fee, Advertising hay Affiliate model.",
    )

    add_heading(doc, "12.3. Luồng thanh toán VNPAY (tóm tắt)", 2)
    add_bullet(doc, "1. POST /api/orders paymentMethod=vnpay → buildVnpayPaymentUrl (HMAC signed)")
    add_bullet(doc, "2. Frontend redirect khách sang sandbox.vnpayment.vn")
    add_bullet(doc, "3. VNPAY gọi GET /api/payments/vnpay/return (browser) và /ipn (server)")
    add_bullet(doc, "4. verifyVnpayCallback + processVnpayQuery → paymentStatus=paid")
    add_bullet(doc, "5. Redirect khách về /checkout/vnpay-result?success=1")

    add_heading(doc, "12.4. Kết luận", 2)
    add_body(
        doc,
        "TechPhone là dự án TMĐT full-stack hiện đại sử dụng React 19, Express 5, MongoDB Atlas, "
        "tích hợp VNPAY, GHN, Google Gemini AI và Socket.io realtime chat. Hệ thống áp dụng các thuật toán "
        "bảo mật chuẩn (PBKDF2, JWT HS256, HMAC-SHA512), transaction MongoDB chống oversell, state machine "
        "quản lý đơn hàng và function calling AI tra cứu dữ liệu thật từ database.",
    )

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
